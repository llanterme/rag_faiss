"""Enhanced document ingestion module with advanced extraction capabilities."""

import re
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional
from datetime import datetime
import logging

from langchain_core.documents import Document
from langchain.text_splitter import (
    RecursiveCharacterTextSplitter,
    MarkdownHeaderTextSplitter,
    PythonCodeTextSplitter,
)

logger = logging.getLogger(__name__)

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logger.warning("pdfplumber not available. Table extraction will be limited.")

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF not available. OCR capabilities will be limited.")

try:
    # Import unstructured only when needed to avoid torch conflicts
    UNSTRUCTURED_AVAILABLE = True
    partition_pdf = None
    partition_docx = None
    def _lazy_import_unstructured():
        global partition_pdf, partition_docx
        if partition_pdf is None:
            from unstructured.partition.pdf import partition_pdf
            from unstructured.partition.docx import partition_docx
        return partition_pdf, partition_docx
except ImportError:
    UNSTRUCTURED_AVAILABLE = False
    partition_pdf = None
    partition_docx = None
    def _lazy_import_unstructured():
        raise ImportError("unstructured not available")
    logger.warning("unstructured not available. Advanced layout analysis will be limited.")

import pypdf
from docx import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph


class EnhancedDocumentProcessor:
    """Enhanced document processor with advanced extraction capabilities."""
    
    def __init__(
        self,
        chunk_size: int = 1500,
        chunk_overlap: int = 300,
        extract_tables: bool = True,
        extract_images: bool = False,
        preserve_formatting: bool = True,
        use_ocr: bool = False,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.extract_tables = extract_tables
        self.extract_images = extract_images
        self.preserve_formatting = preserve_formatting
        self.use_ocr = use_ocr
        
        # Initialize text splitters
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ".", "!", "?", ";", ":", " ", ""],
            length_function=len,
        )
        
        self.markdown_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=[
                ("#", "Header 1"),
                ("##", "Header 2"),
                ("###", "Header 3"),
            ]
        )
    
    def process_document(self, file_path: Path) -> List[Document]:
        """Process a document and return list of Document objects with metadata."""
        if file_path.suffix.lower() == ".pdf":
            return self.process_pdf(file_path)
        elif file_path.suffix.lower() in [".docx", ".doc"]:
            return self.process_docx(file_path)
        elif file_path.suffix.lower() in [".txt", ".md"]:
            return self.process_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
    
    def process_pdf(self, path: Path) -> List[Document]:
        """Process PDF with multiple extraction methods based on availability."""
        documents = []
        
        # Try advanced extraction first if available (skip in UI mode to avoid torch conflicts)
        from src.config import settings
        if (UNSTRUCTURED_AVAILABLE and 
            (self.extract_tables or self.preserve_formatting) and 
            not settings.disable_unstructured_ui):
            try:
                documents = self._process_pdf_unstructured(path)
                if documents:
                    return documents
            except Exception as e:
                logger.warning(f"Unstructured PDF processing failed: {e}")
        
        if PDFPLUMBER_AVAILABLE and self.extract_tables:
            try:
                documents = self._process_pdf_pdfplumber(path)
                if documents:
                    return documents
            except Exception as e:
                logger.warning(f"PDFPlumber processing failed: {e}")
        
        if PYMUPDF_AVAILABLE and self.use_ocr:
            try:
                documents = self._process_pdf_pymupdf(path)
                if documents:
                    return documents
            except Exception as e:
                logger.warning(f"PyMuPDF processing failed: {e}")
        
        # Fallback to basic PyPDF
        return self._process_pdf_pypdf(path)
    
    def _process_pdf_unstructured(self, path: Path) -> List[Document]:
        """Process PDF using unstructured library for advanced layout analysis."""
        partition_pdf_fn, _ = _lazy_import_unstructured()
        elements = partition_pdf_fn(
            filename=str(path),
            strategy="hi_res" if self.preserve_formatting else "fast",
            infer_table_structure=self.extract_tables,
            extract_images_in_pdf=self.extract_images,
        )
        
        documents = []
        current_section = ""
        current_text = []
        
        for element in elements:
            metadata = {
                "source": str(path),
                "file_type": "pdf",
                "element_type": element.category,
                "page_number": element.metadata.page_number if hasattr(element.metadata, 'page_number') else None,
            }
            
            if element.category in ["Title", "Header"]:
                # Start new section
                if current_text:
                    doc = Document(
                        page_content="\n".join(current_text),
                        metadata={**metadata, "section": current_section}
                    )
                    documents.append(doc)
                    current_text = []
                current_section = str(element)
            
            current_text.append(str(element))
        
        # Add final section
        if current_text:
            doc = Document(
                page_content="\n".join(current_text),
                metadata={
                    "source": str(path),
                    "file_type": "pdf",
                    "section": current_section
                }
            )
            documents.append(doc)
        
        return self._split_documents(documents)
    
    def _process_pdf_pdfplumber(self, path: Path) -> List[Document]:
        """Process PDF using pdfplumber for table extraction."""
        documents = []
        
        with pdfplumber.open(path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text
                text = page.extract_text() or ""
                
                # Extract tables
                tables = page.extract_tables() if self.extract_tables else []
                
                metadata = {
                    "source": str(path),
                    "file_type": "pdf",
                    "page_number": page_num,
                    "total_pages": len(pdf.pages),
                }
                
                # Add text content
                if text.strip():
                    doc = Document(page_content=text, metadata=metadata)
                    documents.append(doc)
                
                # Add tables as separate documents
                for i, table in enumerate(tables):
                    table_text = self._format_table(table)
                    table_metadata = {
                        **metadata,
                        "element_type": "table",
                        "table_index": i
                    }
                    doc = Document(page_content=table_text, metadata=table_metadata)
                    documents.append(doc)
        
        return self._split_documents(documents)
    
    def _process_pdf_pymupdf(self, path: Path) -> List[Document]:
        """Process PDF using PyMuPDF with OCR support."""
        documents = []
        doc = fitz.open(path)
        
        for page_num, page in enumerate(doc, 1):
            # Extract text with layout preservation
            text = page.get_text("text" if not self.preserve_formatting else "blocks")
            
            if isinstance(text, list):  # blocks mode returns list
                text = "\n".join([block[4] for block in text if block[6] == 0])
            
            # OCR if text is minimal
            if self.use_ocr and len(text.strip()) < 100:
                pix = page.get_pixmap()
                text = page.get_text("text")  # Re-extract after rendering
            
            metadata = {
                "source": str(path),
                "file_type": "pdf",
                "page_number": page_num,
                "total_pages": len(doc),
                "author": doc.metadata.get("author", ""),
                "title": doc.metadata.get("title", ""),
                "creation_date": doc.metadata.get("creationDate", ""),
            }
            
            if text.strip():
                doc_obj = Document(page_content=text, metadata=metadata)
                documents.append(doc_obj)
        
        doc.close()
        return self._split_documents(documents)
    
    def _process_pdf_pypdf(self, path: Path) -> List[Document]:
        """Basic PDF processing using PyPDF."""
        reader = pypdf.PdfReader(path)
        documents = []
        
        # Extract metadata
        pdf_metadata = reader.metadata or {}
        
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            
            metadata = {
                "source": str(path),
                "file_type": "pdf",
                "page_number": page_num,
                "total_pages": len(reader.pages),
                "author": pdf_metadata.get("/Author", ""),
                "title": pdf_metadata.get("/Title", ""),
                "creation_date": str(pdf_metadata.get("/CreationDate", "")),
            }
            
            if text.strip():
                doc = Document(page_content=text, metadata=metadata)
                documents.append(doc)
        
        return self._split_documents(documents)
    
    def process_docx(self, path: Path) -> List[Document]:
        """Process DOCX with enhanced extraction."""
        from src.config import settings
        if (UNSTRUCTURED_AVAILABLE and 
            self.preserve_formatting and 
            not settings.disable_unstructured_ui):
            try:
                return self._process_docx_unstructured(path)
            except Exception as e:
                logger.warning(f"Unstructured DOCX processing failed: {e}")
        
        return self._process_docx_python_docx(path)
    
    def _process_docx_unstructured(self, path: Path) -> List[Document]:
        """Process DOCX using unstructured library."""
        _, partition_docx_fn = _lazy_import_unstructured()
        elements = partition_docx_fn(
            filename=str(path),
            infer_table_structure=self.extract_tables,
        )
        
        documents = []
        current_section = ""
        
        for element in elements:
            metadata = {
                "source": str(path),
                "file_type": "docx",
                "element_type": element.category,
            }
            
            if element.category in ["Title", "Header"]:
                current_section = str(element)
            
            doc = Document(
                page_content=str(element),
                metadata={**metadata, "section": current_section}
            )
            documents.append(doc)
        
        return self._split_documents(documents)
    
    def _process_docx_python_docx(self, path: Path) -> List[Document]:
        """Enhanced DOCX processing using python-docx."""
        doc = DocxDocument(path)
        documents = []
        current_section = ""
        
        # Extract document properties
        core_props = doc.core_properties
        doc_metadata = {
            "source": str(path),
            "file_type": "docx",
            "author": core_props.author or "",
            "title": core_props.title or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
        }
        
        # Process paragraphs and tables
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                para = Paragraph(element, doc)
                if para.style.name.startswith('Heading'):
                    current_section = para.text
                
                if para.text.strip():
                    metadata = {
                        **doc_metadata,
                        "section": current_section,
                        "style": para.style.name,
                    }
                    documents.append(
                        Document(page_content=para.text, metadata=metadata)
                    )
            
            elif element.tag.endswith('tbl') and self.extract_tables:  # Table
                table = Table(element, doc)
                table_text = self._extract_docx_table(table)
                if table_text.strip():
                    metadata = {
                        **doc_metadata,
                        "section": current_section,
                        "element_type": "table",
                    }
                    documents.append(
                        Document(page_content=table_text, metadata=metadata)
                    )
        
        return self._split_documents(documents)
    
    def _extract_docx_table(self, table: Table) -> str:
        """Extract table content from DOCX."""
        rows = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(row_data))
        return "\n".join(rows)
    
    def process_text(self, path: Path) -> List[Document]:
        """Process text or markdown files."""
        with open(path, "r", encoding="utf-8") as file:
            content = file.read()
        
        metadata = {
            "source": str(path),
            "file_type": path.suffix.lower()[1:],  # Remove the dot
            "file_size": path.stat().st_size,
            "modified": datetime.fromtimestamp(path.stat().st_mtime).isoformat(),
        }
        
        # Use markdown splitter for .md files
        if path.suffix.lower() == ".md" and self.preserve_formatting:
            md_docs = self.markdown_splitter.split_text(content)
            for doc in md_docs:
                doc.metadata.update(metadata)
            return self._split_documents(md_docs)
        
        # Regular text processing
        doc = Document(page_content=content, metadata=metadata)
        return self._split_documents([doc])
    
    def _split_documents(self, documents: List[Document]) -> List[Document]:
        """Split documents into chunks while preserving metadata."""
        split_docs = []
        
        for doc in documents:
            # Skip splitting for tables and short content
            if (doc.metadata.get("element_type") == "table" or 
                len(doc.page_content) <= self.chunk_size):
                split_docs.append(doc)
                continue
            
            # Split long documents
            chunks = self.text_splitter.split_text(doc.page_content)
            for i, chunk in enumerate(chunks):
                chunk_metadata = {
                    **doc.metadata,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                }
                split_docs.append(
                    Document(page_content=chunk, metadata=chunk_metadata)
                )
        
        return split_docs
    
    def _format_table(self, table: List[List[Any]]) -> str:
        """Format table data as markdown-style text."""
        if not table:
            return ""
        
        formatted_rows = []
        for row in table:
            formatted_row = " | ".join(str(cell) if cell else "" for cell in row)
            formatted_rows.append(formatted_row)
        
        # Add header separator after first row
        if len(formatted_rows) > 1:
            header_sep = " | ".join(["---"] * len(table[0]))
            formatted_rows.insert(1, header_sep)
        
        return "\n".join(formatted_rows)
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Advanced text cleaning while preserving structure."""
        # Remove excessive whitespace while preserving paragraph breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove common headers/footers patterns
        patterns = [
            r'^\s*[Pp]age\s+\d+\s*(?:of\s+\d+)?\s*$',
            r'^\s*\d+\s*$',  # Page numbers alone
            r'^\s*[Cc]onfidential\s*$',
            r'^\s*[Dd]raft\s*$',
        ]
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Skip lines matching header/footer patterns
            if not any(re.match(pattern, line) for pattern in patterns):
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines).strip()